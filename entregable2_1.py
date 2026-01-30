#FR
import pandas as pd

# Cargar los archivos CSV
ruta_pbi = r"C:\Users\Fidel\Desktop\UNEMI POSGRADO\ANALISIS CUANTITATIVO AVANZADO\Unidad_2\entregable 2\pib.csv"
ruta_gini = r"C:\Users\Fidel\Desktop\UNEMI POSGRADO\ANALISIS CUANTITATIVO AVANZADO\Unidad_2\entregable 2\gini.csv"


#Aquí convierto mis .csv en dataframes
pbi = pd.read_csv(ruta_pbi, skiprows=4)
gini = pd.read_csv(ruta_gini, skiprows=4)

# Eliminar la última columna (que aparece como Unnamed: 69)
#Las bases de datos originales del banco mundial, tienen una columna vacia
#entinces elimino la ultima columna para que no me de problemas al usar merge

pbi = pbi.iloc[:, :-1]
gini = gini.iloc[:, :-1]

# Elimino las columnas innecesarias
pbi = pbi.drop(columns=["Indicator Name", "Indicator Code"])
gini = gini.drop(columns=["Indicator Name", "Indicator Code"])

# Transformar a formato long (es mi reshape long en stata)
pbi_long = pbi.melt(id_vars=["Country Name","Country Code"], 
                    var_name="anio", value_name="pbi_crecimiento")
gini_long = gini.melt(id_vars=["Country Name","Country Code"], 
                      var_name="anio", value_name="gini")

#exporto el resultado de los dataframes ya transformados para darles un vistazo rapido
#pbi_long.to_excel("pbi_long.xlsx", index=False)
#gini_long.to_excel("gini_long.xlsx", index=False)

#Unifico las tablas con merge
df = pd.merge(pbi_long, gini_long,
              on=["Country Name","Country Code","anio"],
              how="inner")
"""
#Verifico cuantas filas se unieron correctamente de acuerdo a las condiciones de
# pasis, codigo y anio

Este es el outlet del terminal
_merge
both          17290
left_only         0
right_only        0
"""

df = pd.merge(pbi_long, gini_long,
              on=["Country Name","Country Code","anio"],
              how="outer",   # usa outer para ver todo
              indicator=True)

print(df["_merge"].value_counts())

df.to_excel("semitotal.xlsx", index=False)
""""
Convertir año a número entero, hago esto por que al unificar con merge,
la variable la variable anio queda como str y no me sirve par analisis, entonces la conviero a numero entero
"""
df["anio"] = df["anio"].astype(int)

print(df.head())
#Exporto a excel para verificar el dataframe unificado
#df.to_excel("unificado.xlsx", index=False)



#Agrego el metadata que contiene region, tipo de ingreso, etc
ruta_meta = r"C:\Users\Fidel\Desktop\UNEMI POSGRADO\ANALISIS CUANTITATIVO AVANZADO\Unidad_2\entregable 2\API_NY.GDP.MKTP.KD.ZG_DS2_es_csv_v2_277\meta.csv"

meta = pd.read_csv(ruta_meta)

#Los metadatos tienen una columna unnamed4, la elimino a la maldita para que no me de problemas 
meta = meta.drop(columns=["Unnamed: 4"])

"""
el archivo original de los metadatos tiene errores valores de ingreso.* aparecen en
country name, ademas elimino los agregados que no me sirven para analisis
"""

# 1. Eliminar filas donde Country Name contiene "Ingreso"
meta = meta[~meta["Country Name"].str.contains("Ingreso", na=False)]

# 2. Eliminar filas donde Income_Group es "Agregados" o "No clasificado"
meta = meta[(meta["Income_Group"] != "Agregados") & 
            (meta["Income_Group"] != "No clasificado")]
#Exporto a excel para revisar
meta.to_excel("meta.xlsx", index=False)

#Uno mi metadata 
df = pd.merge(df, meta[["Country Code","Region","Income_Group"]],
              on="Country Code", how="inner")

print("Se unieron correctamente:", df.shape[0], "filas")

#Resultado del terminal: Se unieron correctamente: 13975 filas

#Ordeno las variables para que quede bien bonito (de acuerdo a la consigna)
df = df[[
    "anio",
    "Country Name",
    "Country Code",
    "Region",
    "Income_Group",
    "pbi_crecimiento",
    "gini"
]]

print(df.head())

#Reviso mi dataframe
#df.to_excel("total.xlsx", index=False)

#Agrupo las variables por tipo de ingreso
mapa = {
    "Ingreso alto": "Ingreso alto",
    "Ingreso mediano alto": "Ingreso medio",
    "Países de ingreso mediano bajo": "Ingreso medio",
    "Países de ingreso bajo": "Ingreso bajo"
}

df["Ingreso_cat"] = df["Income_Group"].map(mapa)

#Creo la columna numérica de ingreso
codigos = {
    "Ingreso alto": 1,
    "Ingreso medio": 2,
    "Ingreso bajo": 3
}
df["Ingreso_num"] = df["Ingreso_cat"].map(codigos)

#Reviso mi dataframe final
df.to_excel("total_tipo.xlsx", index=False)


# 1. Contar observaciones por país
conteo = df.groupby("Country Name")["anio"].nunique().reset_index()
conteo = conteo.rename(columns={"anio":"n_obs"})

# 2. Filtrar países con al menos 30 observaciones
paises_validos = conteo[conteo["n_obs"] >= 30]["Country Name"]

# 3. Quedarse solo con esos países en el DataFrame principal
df_filtrado = df[df["Country Name"].isin(paises_validos)]

# 4. Verificar intervalo temporal homogéneo
anio_min = df_filtrado.groupby("Country Name")["anio"].min().max()  # el máximo de los mínimos
anio_max = df_filtrado.groupby("Country Name")["anio"].max().min()  # el mínimo de los máximos

print("Intervalo homogéneo común:", anio_min, "-", anio_max)

# 5. Delimitar la muestra al intervalo temporal común
df_final = df_filtrado[(df_filtrado["anio"] >= anio_min) & (df_filtrado["anio"] <= anio_max)]

# 6. Revisar resultado
print("Número de países seleccionados:", df_final["Country Name"].nunique())
print("Número de observaciones totales:", df_final.shape[0])
print(df_final.head())





# 1. Eliminar filas incompletas (sin PIB o sin Gini)
df_final = df_final.dropna(subset=["pbi_crecimiento", "gini"])

# 2. Contar observaciones válidas por país
conteo = df_final.groupby("Country Name")["anio"].nunique().reset_index()
conteo = conteo.rename(columns={"anio":"n_obs"})

# 3. Filtrar países con al menos 30 observaciones completas
paises_validos = conteo[conteo["n_obs"] >= 30]["Country Name"]
df_final = df_final[df_final["Country Name"].isin(paises_validos)]

# 4. Verificar intervalo temporal homogéneo
anio_min = df_final.groupby("Country Name")["anio"].min().max()  # el máximo de los mínimos
anio_max = df_final.groupby("Country Name")["anio"].max().min()  # el mínimo de los máximos
print("Intervalo homogéneo común:", anio_min, "-", anio_max)

# 5. Delimitar la muestra al intervalo temporal común
df_final = df_final[(df_final["anio"] >= anio_min) & (df_final["anio"] <= anio_max)]

# 6. Revisar resultado
print("Número de países seleccionados:", df_final["Country Name"].nunique())
print("Número de observaciones totales:", df_final.shape[0])
print(df_final.head())

#Reviso mi dataframe final final final (3:00 AM)
df_final.to_excel("df_final.xlsx", index=False)


#=====HASTA AQUI SE CUMPLE LOS PUNTOS DEL 1 AL 3
# incluye: carga de las 2 DBA's 
# Unión de las 2 DBA
# Integracion con los metadatos del banco mundial
# filtros de acuerdo a la consigna del item 3)

#==================a
# Crear tabla cruzada: filas = regiones, columnas = niveles de ingreso
tabla_cruzada = pd.crosstab(
    index=df_final["Region"],          # filas = regiones
    columns=df_final["Ingreso_cat"],   # columnas = nivel de ingreso
    values=df_final["Country Name"],   # valores = países
    aggfunc=lambda x: x.nunique(),     # contar países únicos
    dropna=False
)

# Mostrar la tabla en consola
print(tabla_cruzada)

# Exportar a Excel
tabla_cruzada.to_excel("tabla_cruzada.xlsx")

#===================b

import numpy as np

# Definir función para calcular todas las métricas
def resumen_region(x):
    return pd.Series({
        "N": x.count(),
        "Media": x.mean(),
        "Mediana": x.median(),
        "Desv.Std": x.std(),
        "CV": x.std() / x.mean() if x.mean() != 0 else np.nan,
        "Mínimo": x.min(),
        "Máximo": x.max()
    })

# Aplicar por región y variable elegida
resumen_pbi = df_final.groupby("Region")["pbi_crecimiento"].apply(resumen_region)
resumen_gini = df_final.groupby("Region")["gini"].apply(resumen_region)

# Unir ambos resultados en una sola tabla
resumen_total = pd.concat([resumen_pbi, resumen_gini], axis=1, keys=["PIB", "Gini"])

print(resumen_total)

# Exportar a Excel
resumen_total.to_excel("resumen_regiones.xlsx")

#===============c

import numpy as np

# Definir función para calcular todas las métricas
def resumen_ingreso(x):
    return pd.Series({
        "N": x.count(),
        "Media": x.mean(),
        "Mediana": x.median(),
        "Desv.Std": x.std(),
        "CV": x.std() / x.mean() if x.mean() != 0 else np.nan,
        "Mínimo": x.min(),
        "Máximo": x.max()
    })

# Aplicar por tipo de ingreso y variable elegida
resumen_pbi = df_final.groupby("Ingreso_cat")["pbi_crecimiento"].apply(resumen_ingreso)
resumen_gini = df_final.groupby("Ingreso_cat")["gini"].apply(resumen_ingreso)

# Unir ambos resultados en una sola tabla
resumen_total = pd.concat([resumen_pbi, resumen_gini], axis=1, keys=["PIB", "Gini"])

print(resumen_total)

# Exportar a Excel
resumen_total.to_excel("resumen_ingreso.xlsx")

#==============d

# Calcular promedios por región y nivel de ingreso
tabla_promedios = df_final.groupby(["Region", "Ingreso_cat"])[["pbi_crecimiento", "gini"]].mean().reset_index()

print(tabla_promedios)

# Exportar a Excel
tabla_promedios.to_excel("promedios_region_ingreso.xlsx", index=False)

# Identificar mayor y menor promedio en cada indicador
max_pbi = tabla_promedios.loc[tabla_promedios["pbi_crecimiento"].idxmax()]
min_pbi = tabla_promedios.loc[tabla_promedios["pbi_crecimiento"].idxmin()]

max_gini = tabla_promedios.loc[tabla_promedios["gini"].idxmax()]
min_gini = tabla_promedios.loc[tabla_promedios["gini"].idxmin()]

print("\nMayor promedio PIB:", max_pbi["Region"], "-", max_pbi["Ingreso_cat"], "=", round(max_pbi["pbi_crecimiento"],2))
print("Menor promedio PIB:", min_pbi["Region"], "-", min_pbi["Ingreso_cat"], "=", round(min_pbi["pbi_crecimiento"],2))

print("\nMayor promedio Gini:", max_gini["Region"], "-", max_gini["Ingreso_cat"], "=", round(max_gini["gini"],2))
print("Menor promedio Gini:", min_gini["Region"], "-", min_gini["Ingreso_cat"], "=", round(min_gini["gini"],2))

#================e

import seaborn as sns
import matplotlib.pyplot as plt

# 1. Calcular ranking por año
df_final["rank_pbi"] = df_final.groupby("anio")["pbi_crecimiento"].rank(method="first", ascending=False)
df_final["rank_gini"] = df_final.groupby("anio")["gini"].rank(method="first", ascending=False)

# 2. Seleccionar top 10 países según promedio en todo el período
top10_pbi = df_final.groupby("Country Name")["pbi_crecimiento"].mean().nlargest(10).index
df_pbi_top10 = df_final[df_final["Country Name"].isin(top10_pbi)]

top10_gini = df_final.groupby("Country Name")["gini"].mean().nlargest(10).index
df_gini_top10 = df_final[df_final["Country Name"].isin(top10_gini)]

# 3. Gráfico tipo bumpline para PIB
plt.figure(figsize=(10,6))
sns.lineplot(
    data=df_pbi_top10,
    x="anio", y="rank_pbi", hue="Country Name", marker="o"
)
plt.gca().invert_yaxis()  # rank 1 arriba
plt.title("Top 10 países por PIB (1992–2020)")
plt.show()

# 4. Gráfico tipo bumpline para Gini
plt.figure(figsize=(10,6))
sns.lineplot(
    data=df_gini_top10,
    x="anio", y="rank_gini", hue="Country Name", marker="o"
)
plt.gca().invert_yaxis()
plt.title("Top 10 países por Gini (1992–2020)")
plt.show()

#=============f

import seaborn as sns
import matplotlib.pyplot as plt

# 1. Calcular ranking por año
df_final["rank_pbi"] = df_final.groupby("anio")["pbi_crecimiento"].rank(method="first", ascending=False)
df_final["rank_gini"] = df_final.groupby("anio")["gini"].rank(method="first", ascending=False)

# 2. Seleccionar top 10 países según promedio en todo el período
top10_pbi = df_final.groupby("Country Name")["pbi_crecimiento"].mean().nlargest(10).index
top10_gini = df_final.groupby("Country Name")["gini"].mean().nlargest(10).index

# 3. Filtrar con "select()" equivalente en Python → .isin()
df_pbi_select = df_final[df_final["Country Name"].isin(top10_pbi)]
df_gini_select = df_final[df_final["Country Name"].isin(top10_gini)]

# 4. Gráfico tipo bumpline para PIB
plt.figure(figsize=(10,6))
sns.lineplot(
    data=df_pbi_select,
    x="anio", y="rank_pbi", hue="Country Name", marker="o"
)
plt.gca().invert_yaxis()  # rank 1 arriba
plt.title("Trayectoria continua del Top 10 PIB (1992–2020)")
plt.show()

# 5. Gráfico tipo bumpline para Gini
plt.figure(figsize=(10,6))
sns.lineplot(
    data=df_gini_select,
    x="anio", y="rank_gini", hue="Country Name", marker="o"
)
plt.gca().invert_yaxis()
plt.title("Trayectoria continua del Top 10 Gini (1992–2020)")
plt.show()

#================g

# Colapsar: promedio de cada indicador por país
df_collapse = df_final.groupby(
    ["Country Name", "Region", "Ingreso_cat"]
)[["pbi_crecimiento", "gini"]].mean().reset_index()

print(df_collapse.head())

# Resumen detallado estilo summarize, detail
resumen_pbi = df_collapse["pbi_crecimiento"].describe(percentiles=[.25,.5,.75,.9,.95])
resumen_gini = df_collapse["gini"].describe(percentiles=[.25,.5,.75,.9,.95])

print("\nResumen PIB:\n", resumen_pbi)
print("\nResumen Gini:\n", resumen_gini)

# Convertir los resúmenes a DataFrame y exportar
resumen_df = pd.DataFrame({
    "PIB": resumen_pbi,
    "Gini": resumen_gini
})

resumen_df.to_excel("resumen_detail.xlsx")


import seaborn as sns
import matplotlib.pyplot as plt

# Histograma PIB
plt.figure(figsize=(8,5))
sns.histplot(df_collapse["pbi_crecimiento"], bins=20, kde=True)
plt.title("Histograma del promedio de crecimiento del PIB")
plt.xlabel("PIB promedio (%)")
plt.show()

# Histograma Gini
plt.figure(figsize=(8,5))
sns.histplot(df_collapse["gini"], bins=20, kde=True)
plt.title("Histograma del promedio del índice de Gini")
plt.xlabel("Gini promedio")
plt.show()

#=======================h
"""
import seaborn as sns
import matplotlib.pyplot as plt

# =========================
# Histograma por tipo de ingreso (3 subplots)
# =========================
g = sns.FacetGrid(df_final, col="Ingreso_cat", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="pbi_crecimiento", bins=20, kde=True)
g.set_axis_labels("PIB promedio (%)", "Frecuencia")
g.set_titles(col_template="{col_name}")
plt.subplots_adjust(top=0.8)
g.fig.suptitle("Distribución del PIB por tipo de ingreso")
plt.show()

g = sns.FacetGrid(df_final, col="Ingreso_cat", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="gini", bins=20, kde=True)
g.set_axis_labels("Gini promedio", "Frecuencia")
g.set_titles(col_template="{col_name}")
plt.subplots_adjust(top=0.8)
g.fig.suptitle("Distribución del Gini por tipo de ingreso")
plt.show()

# =========================
# Histograma por región (7 subplots)
# =========================
g = sns.FacetGrid(df_final, col="Region", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="pbi_crecimiento", bins=20, kde=True)
g.set_axis_labels("PIB promedio (%)", "Frecuencia")
g.set_titles(col_template="{col_name}")
plt.subplots_adjust(top=0.9)
g.fig.suptitle("Distribución del PIB por región")
plt.show()

g = sns.FacetGrid(df_final, col="Region", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="gini", bins=20, kde=True)
g.set_axis_labels("Gini promedio", "Frecuencia")
g.set_titles(col_template="{col_name}")
plt.subplots_adjust(top=0.9)
g.fig.suptitle("Distribución del Gini por región")
plt.show()

"""
import seaborn as sns
import matplotlib.pyplot as plt

# =========================
# Histograma por tipo de ingreso
# =========================
g = sns.FacetGrid(df_final, col="Ingreso_cat", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="pbi_crecimiento", bins=20, kde=True)
g.set_axis_labels("PIB promedio (%)", "Frecuencia")
g.set_titles(col_template="{col_name}")
g.fig.suptitle("Distribución del PIB por tipo de ingreso", fontsize=14)
plt.subplots_adjust(top=0.85)
plt.savefig("histograma_ingreso_pbi.png", dpi=300, bbox_inches="tight")
plt.show()

g = sns.FacetGrid(df_final, col="Ingreso_cat", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="gini", bins=20, kde=True)
g.set_axis_labels("Gini promedio", "Frecuencia")
g.set_titles(col_template="{col_name}")
g.fig.suptitle("Distribución del Gini por tipo de ingreso", fontsize=14)
plt.subplots_adjust(top=0.85)
plt.savefig("histograma_ingreso_gini.png", dpi=300, bbox_inches="tight")
plt.show()

# =========================
# Histograma por región
# =========================
g = sns.FacetGrid(df_final, col="Region", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="pbi_crecimiento", bins=20, kde=True)
g.set_axis_labels("PIB promedio (%)", "Frecuencia")
g.set_titles(col_template="{col_name}")
g.fig.suptitle("Distribución del PIB por región", fontsize=14)
plt.subplots_adjust(top=0.88)
plt.savefig("histograma_region_pbi.png", dpi=300, bbox_inches="tight")
plt.show()

g = sns.FacetGrid(df_final, col="Region", col_wrap=3, sharex=True, sharey=True)
g.map_dataframe(sns.histplot, x="gini", bins=20, kde=True)
g.set_axis_labels("Gini promedio", "Frecuencia")
g.set_titles(col_template="{col_name}")
g.fig.suptitle("Distribución del Gini por región", fontsize=14)
plt.subplots_adjust(top=0.88)
plt.savefig("histograma_region_gini.png", dpi=300, bbox_inches="tight")
plt.show()

#===============i
import pandas as pd
from docx import Document

# Reemplazar valores NaN en Región
# Esto lo hago por que gropuby, no me estaba tomando en cuenta regiones vacias
df_final["Region"] = df_final["Region"].fillna("Sin región definida")

# Agrupar por país, región y tipo de ingreso
reporte = df_final.groupby(
    ["Country Name", "Region", "Ingreso_cat"], as_index=False
).agg({
    "pbi_crecimiento": "mean",
    "gini": "mean"
})

# Ordenar por PIB promedio de mayor a menor
reporte = reporte.sort_values(by="pbi_crecimiento", ascending=False)

# Renombrar columnas
reporte = reporte.rename(columns={
    "Country Name": "País",
    "Region": "Región",
    "Ingreso_cat": "Tipo de ingreso",
    "pbi_crecimiento": "PIB promedio (%)",
    "gini": "Gini promedio"
})

# Crear documento Word
doc = Document()
doc.add_heading("Reporte de Promedios por País", level=1)

# Crear tabla
tabla = doc.add_table(rows=1, cols=5)
tabla.style = "Table Grid"

# Encabezados
hdr_cells = tabla.rows[0].cells
hdr_cells[0].text = "País"
hdr_cells[1].text = "Región"
hdr_cells[2].text = "Tipo de ingreso"
hdr_cells[3].text = "PIB promedio (%)"
hdr_cells[4].text = "Gini promedio"

# Llenar filas
for _, fila in reporte.iterrows():
    row_cells = tabla.add_row().cells
    row_cells[0].text = str(fila["País"])
    row_cells[1].text = str(fila["Región"])
    row_cells[2].text = str(fila["Tipo de ingreso"])
    row_cells[3].text = f"{fila['PIB promedio (%)']:.2f}"
    row_cells[4].text = f"{fila['Gini promedio']:.2f}"

# Guardar documento
doc.save("reporte_promedios.docx")

